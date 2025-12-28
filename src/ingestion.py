import os
import re
import logging
from typing import List, Dict, Any
from pathlib import Path

# Try importing libraries, handle errors if dependencies are missing
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

from src.interfaces import Document

# Setup logger for this module
logger = logging.getLogger(__name__)

class TextCleaner:
    """
    Responsible for normalizing text and removing noise.
    """
    
    @staticmethod
    def clean(text: str) -> str:
        if not text:
            return ""
            
        # 1. Normalize whitespace (replace tabs, newlines, multi-spaces with single space)
        text = re.sub(r'\s+', ' ', text).strip()
        
        # 2. Remove non-informative text (e.g., generic headers/footers)
        text = re.sub(r'Page \d+ of \d+', '', text)
        
        # 3. Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable())
        
        return text

class IngestionEngine:
    """
    Factory class to load different file types.
    Enforces validation on loaded documents.
    """

    def load_file(self, file_path: str) -> List[Document]:
        """
        Loads a file and returns a list of Document objects.
        Raises ValueError if the file yields no content.
        """
        path = Path(file_path)
        
        # Guard Clause: File existence
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()
        docs = []

        try:
            if extension == '.pdf':
                docs = self._load_pdf(path)
            elif extension == '.docx':
                docs = self._load_docx(path)
            elif extension == '.txt':
                docs = self._load_txt(path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
        except Exception as e:
            logger.error(f"Failed to load {file_path}: {e}")
            raise e

        # Guard Clause: Empty Result
        if not docs:
            raise ValueError(f"File {file_path} was loaded but contained no valid text content.")

        # Validate Metadata for all loaded docs
        for doc in docs:
            self._validate_document(doc)

        return docs

    def _validate_document(self, doc: Document) -> None:
        """
        Ensures strict contract compliance for documents.
        """
        if not isinstance(doc, Document):
            raise TypeError("Object is not a valid Document instance.")
        
        # Check required metadata keys
        required_keys = {'source', 'type'}
        if not all(key in doc.metadata for key in required_keys):
            raise ValueError(f"Document missing required metadata keys: {required_keys - doc.metadata.keys()}")
            
        # Check content
        if not doc.content or not doc.content.strip():
            logger.warning(f"Document from {doc.metadata.get('source')} has empty content.")

    def _load_pdf(self, path: Path) -> List[Document]:
        if not pypdf:
            raise ImportError("pypdf is required. Run: pip install pypdf")
        
        docs = []
        reader = pypdf.PdfReader(str(path))
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if not text:
                continue
                
            cleaned_text = TextCleaner.clean(text)
            
            # Only add if content is substantive
            if len(cleaned_text) > 10: 
                docs.append(Document(
                    content=cleaned_text,
                    metadata={"source": path.name, "page": i + 1, "type": "pdf"}
                ))
        return docs

    def _load_docx(self, path: Path) -> List[Document]:
        if not docx:
            raise ImportError("python-docx is required. Run: pip install python-docx")
        
        doc = docx.Document(str(path))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        text = "\n".join(full_text)
        cleaned_text = TextCleaner.clean(text)
        
        if len(cleaned_text) < 10:
            return []

        return [Document(
            content=cleaned_text,
            metadata={"source": path.name, "type": "docx"}
        )]

    def _load_txt(self, path: Path) -> List[Document]:
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        cleaned_text = TextCleaner.clean(text)
        
        if len(cleaned_text) < 10:
            return []

        return [Document(
            content=cleaned_text,
            metadata={"source": path.name, "type": "txt"}
        )]
    
class TextChunker:
    """
    Splits documents into smaller chunks with overlap to maintain context.
    """
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """
        Chunks a list of documents.
        Raises ValueError if input list is empty or invalid.
        """
        # Guard Clause: Empty Input
        if not documents:
            raise ValueError("Input document list cannot be empty.")

        chunked_docs = []
        
        for doc in documents:
            # Validate input type
            if not isinstance(doc, Document):
                raise TypeError("Processing failed: Input list contains non-Document objects.")

            text = doc.content
            if not text:
                continue
                
            start = 0
            while start < len(text):
                end = start + self.chunk_size
                chunk_text = text[start:end]
                
                # Create new document for chunk, copying original metadata
                chunk_metadata = doc.metadata.copy()
                chunk_metadata["chunk_index"] = len(chunked_docs)
                
                chunked_docs.append(Document(
                    content=chunk_text,
                    metadata=chunk_metadata
                ))
                
                # Move forward, accounting for overlap
                start += (self.chunk_size - self.chunk_overlap)
        
        if not chunked_docs:
             logger.warning("Chunking process resulted in 0 chunks.")

        return chunked_docs
# import os
# import re
# from typing import List, Dict, Any
# from pathlib import Path

# # Try importing libraries, handle errors if dependencies are missing
# try:
#     import pypdf
# except ImportError:
#     pypdf = None

# try:
#     import docx
# except ImportError:
#     docx = None

# from src.interfaces import Document

# class TextCleaner:
#     """
#     Responsible for normalizing text and removing noise.
#     """
    
#     @staticmethod
#     def clean(text: str) -> str:
#         if not text:
#             return ""
            
#         # 1. Normalize whitespace (replace tabs, newlines, multi-spaces with single space)
#         text = re.sub(r'\s+', ' ', text).strip()
        
#         # 2. Remove non-informative text (e.g., generic headers/footers if recognizable)
#         # This is a basic heuristic; can be expanded with specific regex patterns
#         # Example: Removing page numbers like "Page 1 of 10"
#         text = re.sub(r'Page \d+ of \d+', '', text)
        
#         # 3. Remove non-printable characters
#         text = ''.join(char for char in text if char.isprintable())
        
#         return text

# class IngestionEngine:
#     """
#     Factory class to load different file types.
#     """

#     def load_file(self, file_path: str) -> List[Document]:
#         path = Path(file_path)
#         if not path.exists():
#             raise FileNotFoundError(f"File not found: {file_path}")

#         extension = path.suffix.lower()

#         if extension == '.pdf':
#             return self._load_pdf(path)
#         elif extension == '.docx':
#             return self._load_docx(path)
#         elif extension == '.txt':
#             return self._load_txt(path)
#         else:
#             raise ValueError(f"Unsupported file type: {extension}")

#     def _load_pdf(self, path: Path) -> List[Document]:
#         if not pypdf:
#             raise ImportError("pypdf is required. Run: pip install pypdf")
        
#         docs = []
#         reader = pypdf.PdfReader(str(path))
        
#         for i, page in enumerate(reader.pages):
#             text = page.extract_text()
#             cleaned_text = TextCleaner.clean(text)
            
#             # Only add if content is substantive
#             if len(cleaned_text) > 10: 
#                 docs.append(Document(
#                     content=cleaned_text,
#                     metadata={"source": path.name, "page": i + 1, "type": "pdf"}
#                 ))
#         return docs

#     def _load_docx(self, path: Path) -> List[Document]:
#         if not docx:
#             raise ImportError("python-docx is required. Run: pip install python-docx")
        
#         doc = docx.Document(str(path))
#         full_text = []
#         for para in doc.paragraphs:
#             full_text.append(para.text)
        
#         text = "\n".join(full_text)
#         cleaned_text = TextCleaner.clean(text)
        
#         return [Document(
#             content=cleaned_text,
#             metadata={"source": path.name, "type": "docx"}
#         )]

#     def _load_txt(self, path: Path) -> List[Document]:
#         with open(path, 'r', encoding='utf-8') as f:
#             text = f.read()
        
#         cleaned_text = TextCleaner.clean(text)
        
#         return [Document(
#             content=cleaned_text,
#             metadata={"source": path.name, "type": "txt"}
#         )]
    
# class TextChunker:
#     """
#     Splits documents into smaller chunks with overlap to maintain context.
#     """
#     def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
#         self.chunk_size = chunk_size
#         self.chunk_overlap = chunk_overlap

#     def chunk_documents(self, documents: List[Document]) -> List[Document]:
#         chunked_docs = []
        
#         for doc in documents:
#             text = doc.content
#             if not text:
#                 continue
                
#             # Simple character-based sliding window
#             # (For production, consider using a tokenizer-based splitter like tiktoken)
#             start = 0
#             while start < len(text):
#                 end = start + self.chunk_size
#                 chunk_text = text[start:end]
                
#                 # Create new document for chunk, copying original metadata
#                 chunk_metadata = doc.metadata.copy()
#                 chunk_metadata["chunk_index"] = len(chunked_docs)
                
#                 chunked_docs.append(Document(
#                     content=chunk_text,
#                     metadata=chunk_metadata
#                 ))
                
#                 # Move forward, accounting for overlap
#                 start += (self.chunk_size - self.chunk_overlap)
                
#         return chunked_docs